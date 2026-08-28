"""Local text extraction from candidate-provided documents.

This module intentionally performs extraction only. The AI Profile Agent is
responsible for turning extracted evidence into a structured CandidateProfile.
Scanned-image PDFs may contain no extractable text; the user is warned rather
than silently applying OCR.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from pypdf import PdfReader

from .models import CandidateDocumentBundle, ExtractedDocument

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".md", ".txt"}


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(page for page in pages if page)


def _read_docx(path: Path) -> str:
    document = Document(str(path))
    return "\n".join(p.text for p in document.paragraphs if p.text.strip())


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported candidate file type: {path.suffix}")


def collect_candidate_documents(private_dir: Path) -> CandidateDocumentBundle:
    documents: list[ExtractedDocument] = []
    if not private_dir.exists():
        return CandidateDocumentBundle(documents=[])

    for path in sorted(p for p in private_dir.rglob("*") if p.is_file()):
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            continue
        text = extract_text(path)
        kind = path.parent.name
        documents.append(ExtractedDocument(path=str(path), text=text, kind=kind))
    return CandidateDocumentBundle(documents=documents)
